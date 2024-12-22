import os
from django.conf import settings
from docx import Document
from .models import (
    Student, 
    Olympiads, 
    Tutors, 
    Afterschools, 
    Classes
)


class FormDocument:
    def __init__(self, chosen_class, doc_conf):
        self.chosen_class = chosen_class
        self.persons = doc_conf['persons']
        self.selected_activities = [activity for activity in doc_conf['activities']]
        self.teacher_id = Classes.objects.get(id=self.chosen_class).main_teacher_id

    def get_from_db(self):
        group = Classes.objects.get(id=self.chosen_class)

        if (self.persons == 'all'):
            students = Student.objects.filter(classes_id=self.chosen_class)
        else:
            students = Student.objects.filter(id=self.persons)

        return group, students

    def add_student_data_into_doc(self, document, student):
        document.add_paragraph(f'{student.surname} {student.name} {student.patronymic}')

        for activity in self.selected_activities:
            match activity:
                case 'olympiads':
                    olympiads = Olympiads.objects.filter(student_id=student.id)
                    if (len(olympiads) != 0): 
                        p = document.add_paragraph()
                        p.add_run('Олимпиады:').bold = True
                        for o in olympiads:
                            document.add_paragraph(f'Название: {o.name}')
                            if (o.place): document.add_paragraph(f'Место: {o.place}')
                            if (o.info): document.add_paragraph(f'Информация: {o.info}')

                case 'tutors':
                    tutors = Tutors.objects.filter(student_id=student.id)
                    if (len(tutors) != 0): 
                        p = document.add_paragraph()
                        p.add_run('Репетиторы:').bold = True
                        for t in tutors:
                            document.add_paragraph(f'ФИО: {t.surname} {t.name} {t.patronymic}')
                            if (t.subject): document.add_paragraph(f'Предмет: {t.subject}')
                            if (t.info): document.add_paragraph(f'Информация: {t.info}')

                case 'afterschools':
                    afterschools = Afterschools.objects.filter(student_id=student.id)
                    if (len(afterschools) != 0): 
                        p = document.add_paragraph()
                        p.add_run('Кружки:').bold = True
                        for a in afterschools:
                            document.add_paragraph(f'Название: {a.subject}')
                            if (a.info): document.add_paragraph(f'Информация: {a.info}')

    def create_document(self):

        group, students = self.get_from_db()

        document = Document()
        document.add_heading('Отчёт')
        document.add_paragraph(f'Класс: {group.name}')

        for student in students:
            self.add_student_data_into_doc(document, student)

        document.add_heading('Heading, level 1', level=1)
        document.add_paragraph('Intense quote', style='Intense Quote')

        output_path = os.path.join(settings.MEDIA_ROOT, f"docs\doc{self.teacher_id}.docx")
        document.save(output_path)
        return output_path
    